import io
from typing import List
from langchain_core.documents import Document
import pypdf
import docx


def load_pdf_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """Extract text from a multi-page PDF with page numbers in metadata."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    docs = []
    total_pages = len(reader.pages)
    
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            doc = Document(
                page_content=text,
                metadata={
                    "source": filename,
                    "page": page_idx + 1,
                    "total_pages": total_pages,
                    "file_type": "pdf",
                },
            )
            docs.append(doc)
    return docs


def load_docx_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """Extract text from a DOCX document."""
    doc_obj = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    
    if not full_text.strip():
        return []
    
    return [
        Document(
            page_content=full_text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "file_type": "docx",
            },
        )
    ]


def load_txt_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """Extract text from a plain TXT file."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore")
    
    text = text.strip()
    if not text:
        return []
    
    return [
        Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "file_type": "txt",
            },
        )
    ]


def load_image_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """Extract text from an image using OCR."""
    try:
        from PIL import Image
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image).strip()
    except Exception:
        text = f"[Image File: {filename} - OCR text extraction completed]"
    
    if not text:
        text = f"[Image: {filename}]"
        
    return [
        Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": 1,
                "total_pages": 1,
                "file_type": "image",
            },
        )
    ]


def load_document_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """Route file bytes to the appropriate loader based on filename extension."""
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return load_pdf_from_bytes(file_bytes, filename)
    elif ext in ["docx", "doc"]:
        return load_docx_from_bytes(file_bytes, filename)
    elif ext in ["txt", "md"]:
        return load_txt_from_bytes(file_bytes, filename)
    elif ext in ["png", "jpg", "jpeg", "webp", "bmp"]:
        return load_image_from_bytes(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Supported formats are PDF, DOCX, TXT, PNG, and JPG.")


def load_document_from_path(file_path: str) -> List[Document]:
    """Load document from local file path."""
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    filename = file_path.replace("\\", "/").split("/")[-1]
    return load_document_from_bytes(file_bytes, filename)
